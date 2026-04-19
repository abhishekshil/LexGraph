from __future__ import annotations

import io

from ..core import get_logger
from .parser import ParsedDocument


log = get_logger("parser.docx")


def parse_docx_bytes(data: bytes) -> ParsedDocument:
    try:
        import docx  # type: ignore
    except Exception as e:  # noqa: BLE001
        log.warning("python-docx_missing", error=str(e))
        return ParsedDocument(
            text="", page_offsets=[], mime="application/docx", page_count=0
        )

    try:
        doc = docx.Document(io.BytesIO(data))
    except Exception as e:  # noqa: BLE001
        log.warning("docx.parse_failed", error=str(e))
        return ParsedDocument(
            text="", page_offsets=[], mime="application/docx", page_count=0
        )

    parts: list[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))

    text = "\n".join(p for p in parts if p is not None)
    return ParsedDocument(
        text=text,
        page_offsets=[(0, len(text))],
        mime="application/docx",
        page_count=1,
    )
